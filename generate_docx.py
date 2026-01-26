#!/usr/bin/env python3
"""
Script untuk generate MAKALAH_CSV_QL.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_code_block(doc, code):
    """Add a code block with monospace font."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_table(doc, headers, rows):
    """Add a table with headers and rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
    
    return table

def create_document():
    doc = Document()
    
    # ==================== HALAMAN JUDUL ====================
    title = doc.add_heading('', 0)
    title_run = title.add_run('MAKALAH PROYEK AKHIR')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('MATA KULIAH\n').bold = True
    subtitle.add_run('AUTOMATA DAN TEKNIK KOMPILASI').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    project_title = doc.add_heading('CSV_QL: Mini Query Engine untuk Data Nilai Mahasiswa', level=1)
    project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tema = doc.add_paragraph()
    tema.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tema.add_run('Tema Project: Big Data - Mini Query Engine berbasis SQL').italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Disusun oleh:\n')
    info.add_run('[Nama Anggota Kelompok]\n\n')
    info.add_run('Tahun Akademik 2025/2026')
    
    doc.add_page_break()
    
    # ==================== DAFTAR ISI ====================
    add_heading(doc, 'DAFTAR ISI', 1)
    
    toc_items = [
        ('BAB I', 'Pendahuluan', '1'),
        ('BAB II', 'Deskripsi Tema dan Studi Kasus', '3'),
        ('BAB III', 'Analisis Leksikal (Lexer)', '5'),
        ('BAB IV', 'Analisis Sintaksis (Parser)', '8'),
        ('BAB V', 'Analisis Semantik', '11'),
        ('BAB VI', 'Intermediate Representation (IR)', '13'),
        ('BAB VII', 'Execution Engine', '15'),
        ('BAB VIII', 'Simulasi Automata (DFA)', '17'),
        ('BAB IX', 'Pengujian dan Hasil', '19'),
        ('BAB X', 'Kesimpulan', '22'),
        ('', 'Daftar Pustaka', '23'),
        ('', 'Lampiran', '24'),
    ]
    
    for bab, title, page in toc_items:
        p = doc.add_paragraph()
        if bab:
            p.add_run(f'{bab} - ').bold = True
        p.add_run(f'{title}')
        p.add_run(f'\t\t{page}')
    
    doc.add_page_break()
    
    # ==================== BAB I PENDAHULUAN ====================
    add_heading(doc, 'BAB I - PENDAHULUAN', 1)
    
    add_heading(doc, '1.1 Latar Belakang', 2)
    doc.add_paragraph(
        'Dalam era digital saat ini, data menjadi aset yang sangat penting dalam berbagai bidang, '
        'termasuk pendidikan. Sistem Informasi Akademik (SIMAK/SIAKAD) menghasilkan data dalam jumlah besar, '
        'termasuk data nilai mahasiswa yang umumnya disimpan dalam format CSV (Comma-Separated Values).'
    )
    doc.add_paragraph(
        'Untuk menganalisis data tersebut, pengguna biasanya memerlukan pengetahuan tentang bahasa pemrograman '
        'atau tools khusus. Hal ini menimbulkan hambatan bagi pengguna non-teknis yang ingin melakukan query '
        'sederhana terhadap data.'
    )
    doc.add_paragraph(
        'CSV_QL hadir sebagai solusi berupa Domain Specific Language (DSL) yang memungkinkan pengguna melakukan '
        'query terhadap file CSV menggunakan sintaks mirip SQL yang lebih familiar dan mudah dipahami.'
    )
    
    add_heading(doc, '1.2 Rumusan Masalah', 2)
    problems = [
        'Bagaimana membangun lexical analyzer berbasis DFA untuk tokenisasi query SQL?',
        'Bagaimana merancang parser menggunakan CFG dengan metode recursive descent?',
        'Bagaimana mengimplementasikan analisis semantik untuk validasi query?',
        'Bagaimana merancang Intermediate Representation (IR) untuk eksekusi query?',
        'Bagaimana mengimplementasikan query engine yang dapat mengeksekusi query terhadap file CSV?',
    ]
    for i, problem in enumerate(problems, 1):
        doc.add_paragraph(f'{i}. {problem}')
    
    add_heading(doc, '1.3 Tujuan', 2)
    objectives = [
        'Mengimplementasikan lexical analyzer berbasis DFA',
        'Membangun parser dengan Context-Free Grammar (CFG)',
        'Menyusun Abstract Syntax Tree (AST) dan analisis semantik',
        'Membangun Intermediate Representation (IR) untuk query plan',
        'Mengembangkan execution engine untuk menjalankan query',
    ]
    for i, obj in enumerate(objectives, 1):
        doc.add_paragraph(f'{i}. {obj}')
    
    add_heading(doc, '1.4 Batasan Masalah', 2)
    doc.add_paragraph('• Query terbatas pada operasi SELECT')
    doc.add_paragraph('• Mendukung klausa WHERE dengan operator perbandingan dan logika')
    doc.add_paragraph('• Mendukung klausa LIMIT')
    doc.add_paragraph('• File input harus berformat CSV dengan header')
    
    doc.add_page_break()
    
    # ==================== BAB II DESKRIPSI TEMA ====================
    add_heading(doc, 'BAB II - DESKRIPSI TEMA DAN STUDI KASUS', 1)
    
    add_heading(doc, '2.1 Tema Project', 2)
    doc.add_paragraph(
        'Proyek ini mengambil tema Big Data dengan fokus membangun Mini Query Engine untuk subset SQL.'
    )
    
    # Pengertian Big Data
    add_heading(doc, 'Pengertian Big Data', 3)
    doc.add_paragraph(
        'Big Data adalah istilah yang menggambarkan volume data yang sangat besar, baik terstruktur maupun '
        'tidak terstruktur. Big Data memiliki karakteristik 3V: Volume (ukuran data), Velocity (kecepatan data), '
        'dan Variety (keragaman data). Dalam konteks proyek ini, kami fokus pada pemrosesan data terstruktur '
        'dalam format CSV menggunakan bahasa query.'
    )
    
    # Pengertian Query Engine
    add_heading(doc, 'Pengertian Query Engine', 3)
    doc.add_paragraph(
        'Query Engine adalah komponen software yang bertanggung jawab untuk memproses dan mengeksekusi query '
        'terhadap data. Query engine menerima query dalam bentuk teks, mem-parsing query tersebut, '
        'mengoptimasi rencana eksekusi, dan mengembalikan hasil. Contoh query engine populer adalah '
        'MySQL, PostgreSQL, dan SQLite.'
    )
    
    add_heading(doc, '2.2 Studi Kasus: Query Data Nilai Mahasiswa', 2)
    doc.add_paragraph(
        'Dalam sistem akademik, data nilai mahasiswa sering disimpan dalam format CSV dengan struktur sebagai berikut:'
    )
    
    add_table(doc,
        ['Kolom', 'Tipe', 'Deskripsi'],
        [
            ['nim', 'String', 'Nomor Induk Mahasiswa'],
            ['nama', 'String', 'Nama mahasiswa'],
            ['mata_kuliah', 'String', 'Nama mata kuliah'],
            ['sks', 'Number', 'Jumlah SKS'],
            ['nilai_huruf', 'String', 'Nilai huruf (A, B, C, D, E)'],
            ['nilai_angka', 'Number', 'Nilai angka (0.0 - 4.0)'],
            ['semester', 'Number', 'Semester pengambilan'],
            ['status', 'String', 'Status kelulusan'],
        ]
    )
    
    doc.add_paragraph()
    add_heading(doc, '2.3 Contoh Query', 2)
    add_code_block(doc, '''-- Lihat semua data nilai
SELECT * FROM data_nilai.csv

-- Filter mahasiswa dengan nilai A
SELECT nama, mata_kuliah FROM data_nilai.csv WHERE nilai_huruf = "A"

-- Kombinasi kondisi
SELECT * FROM data_nilai.csv WHERE nilai_angka >= 3.0 AND semester = 5

-- Batasi hasil
SELECT nama, nilai_huruf FROM data_nilai.csv LIMIT 5''')
    
    doc.add_page_break()
    
    # ==================== BAB III ANALISIS LEKSIKAL ====================
    add_heading(doc, 'BAB III - ANALISIS LEKSIKAL (LEXER)', 1)
    
    # Pengertian Lexer
    add_heading(doc, 'Pengertian Lexical Analysis', 3)
    doc.add_paragraph(
        'Lexical Analysis (Analisis Leksikal) adalah tahap pertama dalam proses kompilasi yang bertugas '
        'mengubah aliran karakter (source code) menjadi aliran token. Proses ini juga disebut tokenisasi '
        'atau scanning. Komponen yang melakukan proses ini disebut Lexer atau Scanner.'
    )
    
    add_heading(doc, 'Pengertian Token', 3)
    doc.add_paragraph(
        'Token adalah unit terkecil yang bermakna dalam sebuah bahasa pemrograman. Token terdiri dari '
        'pasangan (tipe_token, nilai). Contoh token: keyword (SELECT, FROM), identifier (nama_kolom), '
        'operator (=, >), dan literal (angka, string).'
    )
    
    add_heading(doc, '3.1 Daftar Token', 2)
    doc.add_paragraph('Lexer CSV_QL mengenali token-token berikut:')
    
    add_table(doc,
        ['Kategori', 'Token', 'Simbol/Pattern'],
        [
            ['Keywords', 'SELECT', 'SELECT'],
            ['', 'FROM', 'FROM'],
            ['', 'WHERE', 'WHERE'],
            ['', 'LIMIT', 'LIMIT'],
            ['', 'AND', 'AND'],
            ['', 'OR', 'OR'],
            ['Operators', 'EQUAL', '='],
            ['', 'NOT_EQUAL', '!= atau <>'],
            ['', 'GREATER_THAN', '>'],
            ['', 'LESS_THAN', '<'],
            ['', 'GREATER_THAN_OR_EQ', '>='],
            ['', 'LESS_THAN_OR_EQ', '<='],
            ['', 'STAR', '*'],
            ['Literals', 'IDENTIFIER', '[a-zA-Z_][a-zA-Z0-9_.]*'],
            ['', 'NUMBER', '[0-9]+(\\.[0-9]+)?'],
            ['', 'STRING_LITERAL', '"[^"]*" atau \'[^\']*\''],
            ['Punctuation', 'COMMA', ','],
        ]
    )
    
    doc.add_paragraph()
    add_heading(doc, '3.2 Regular Expression', 2)
    
    # Pengertian Regex
    add_heading(doc, 'Pengertian Regular Expression', 3)
    doc.add_paragraph(
        'Regular Expression (Regex) adalah notasi formal untuk mendeskripsikan pola string. '
        'Regex digunakan untuk mendefinisikan pola token dalam lexer. Regex dapat dikonversi '
        'menjadi NFA (Non-deterministic Finite Automaton) dan kemudian ke DFA.'
    )
    
    doc.add_paragraph('Berikut adalah regular expression untuk setiap kategori token:')
    add_code_block(doc, '''KEYWORD     := SELECT | FROM | WHERE | LIMIT | AND | OR
IDENTIFIER  := [a-zA-Z_][a-zA-Z0-9_.]*
NUMBER      := [0-9]+(\\.[0-9]+)?
STRING      := "[^"]*" | '[^']*'
OPERATOR    := = | != | <> | > | < | >= | <=
SYMBOL      := * | ,
WHITESPACE  := [ \\t\\n\\r]+''')
    
    add_heading(doc, '3.3 Diagram DFA Lexer', 2)
    
    # Pengertian DFA
    add_heading(doc, 'Pengertian DFA (Deterministic Finite Automaton)', 3)
    doc.add_paragraph(
        'DFA adalah mesin abstrak yang terdiri dari himpunan state terbatas, alphabet input, fungsi transisi, '
        'state awal, dan himpunan state akhir. DFA memproses input karakter per karakter dan berpindah dari '
        'satu state ke state lain berdasarkan fungsi transisi. DFA digunakan dalam lexer untuk mengenali token.'
    )
    
    doc.add_paragraph('Definisi formal DFA: M = (Q, Σ, δ, q0, F)')
    doc.add_paragraph('• Q = {q0, q1, q2, q3, q4, qf} — Himpunan state')
    doc.add_paragraph('• Σ = ASCII characters — Alphabet input')
    doc.add_paragraph('• δ = Fungsi transisi')
    doc.add_paragraph('• q0 = State awal')
    doc.add_paragraph('• F = {qf} — State akhir')
    
    doc.add_paragraph()
    doc.add_paragraph('Deskripsi State:')
    doc.add_paragraph('• q0 (Start): State awal, menunggu input')
    doc.add_paragraph('• q1 (InIdent): Membaca identifier/keyword')
    doc.add_paragraph('• q2 (InNumber): Membaca angka')
    doc.add_paragraph('• q3 (InString): Membaca string literal')
    doc.add_paragraph('• q4 (InOper): Membaca operator')
    doc.add_paragraph('• qf (Accept): State akhir')
    
    add_heading(doc, '3.4 Tabel Transisi DFA', 2)
    add_table(doc,
        ['State', 'Input', 'Next State', 'Action'],
        [
            ['q0', '[a-zA-Z_]', 'q1', 'Start identifier'],
            ['q0', '[0-9]', 'q2', 'Start number'],
            ['q0', '" atau \'', 'q3', 'Start string'],
            ['q0', '=><!=', 'q4', 'Start operator'],
            ['q0', '*,', 'qf', 'Emit symbol'],
            ['q0', 'whitespace', 'q0', 'Skip'],
            ['q1', '[a-zA-Z0-9_.]', 'q1', 'Continue identifier'],
            ['q1', 'otherwise', 'qf', 'Emit IDENTIFIER/KEYWORD'],
            ['q2', '[0-9.]', 'q2', 'Continue number'],
            ['q2', 'otherwise', 'qf', 'Emit NUMBER'],
            ['q3', '[^"\']', 'q3', 'Continue string'],
            ['q3', '" atau \'', 'qf', 'Emit STRING_LITERAL'],
        ]
    )
    
    add_heading(doc, '3.5 Contoh Tokenisasi', 2)
    doc.add_paragraph('Input Query:')
    add_code_block(doc, 'SELECT nama, nilai_huruf FROM data_nilai.csv WHERE status = "Lulus"')
    
    doc.add_paragraph('Output Tokens:')
    add_code_block(doc, '''Token(SELECT)
Token(IDENTIFIER, "nama")
Token(COMMA)
Token(IDENTIFIER, "nilai_huruf")
Token(FROM)
Token(IDENTIFIER, "data_nilai.csv")
Token(WHERE)
Token(IDENTIFIER, "status")
Token(EQUAL)
Token(STRING_LITERAL, "Lulus")''')
    
    doc.add_page_break()
    
    # ==================== BAB IV ANALISIS SINTAKSIS ====================
    add_heading(doc, 'BAB IV - ANALISIS SINTAKSIS (PARSER)', 1)
    
    # Pengertian Parser
    add_heading(doc, 'Pengertian Syntax Analysis', 3)
    doc.add_paragraph(
        'Syntax Analysis (Analisis Sintaksis) adalah tahap kedua dalam kompilasi yang bertugas memeriksa '
        'struktur gramatikal dari aliran token dan membangun parse tree atau AST. Komponen yang melakukan '
        'analisis sintaksis disebut Parser.'
    )
    
    add_heading(doc, '4.1 Context-Free Grammar (CFG)', 2)
    
    # Pengertian CFG
    add_heading(doc, 'Pengertian CFG', 3)
    doc.add_paragraph(
        'Context-Free Grammar (CFG) adalah notasi formal untuk mendefinisikan struktur sintaksis bahasa. '
        'CFG terdiri dari terminal (token), non-terminal (variabel), production rules, dan start symbol. '
        'CFG lebih ekspresif dari Regular Expression dan dapat mendeskripsikan struktur bersarang.'
    )
    
    doc.add_paragraph('Grammar CSV_QL dalam Backus-Naur Form (BNF):')
    add_code_block(doc, '''<query>       ::= <select_stmt>
<select_stmt> ::= SELECT <columns> FROM <table> [<where_clause>] [<limit_clause>]
<columns>     ::= STAR | <column_list>
<column_list> ::= IDENTIFIER (COMMA IDENTIFIER)*
<table>       ::= IDENTIFIER
<where_clause> ::= WHERE <expression>
<limit_clause> ::= LIMIT NUMBER
<expression>  ::= <or_expr>
<or_expr>     ::= <and_expr> (OR <and_expr>)*
<and_expr>    ::= <comparison> (AND <comparison>)*
<comparison>  ::= <leaf> [<comp_op> <leaf>]
<comp_op>     ::= EQUAL | NOT_EQUAL | GREATER_THAN | LESS_THAN | ...
<leaf>        ::= IDENTIFIER | NUMBER | STRING_LITERAL''')
    
    add_heading(doc, '4.2 Operator Precedence', 2)
    add_table(doc,
        ['Level', 'Operator', 'Associativity'],
        [
            ['1 (lowest)', 'OR', 'Left'],
            ['2', 'AND', 'Left'],
            ['3 (highest)', '=, !=, >, <, >=, <=', 'Left'],
        ]
    )
    
    add_heading(doc, '4.3 Metode Parsing: Recursive Descent', 2)
    
    # Pengertian Recursive Descent
    add_heading(doc, 'Pengertian Recursive Descent Parsing', 3)
    doc.add_paragraph(
        'Recursive Descent Parsing adalah teknik parsing top-down di mana setiap non-terminal dalam grammar '
        'diimplementasikan sebagai fungsi rekursif. Parser membaca token dari kiri ke kanan dan membangun '
        'parse tree dari atas ke bawah. Teknik ini mudah diimplementasikan dan dipahami.'
    )
    
    doc.add_paragraph('Setiap non-terminal dalam grammar memiliki fungsi parser yang sesuai:')
    doc.add_paragraph('• parse() → Entry point')
    doc.add_paragraph('• parse_select() → Parse SELECT statement')
    doc.add_paragraph('• parse_columns() → Parse daftar kolom')
    doc.add_paragraph('• parse_expression() → Parse WHERE clause')
    doc.add_paragraph('• parse_logic_or() → Parse operasi OR')
    doc.add_paragraph('• parse_logic_and() → Parse operasi AND')
    doc.add_paragraph('• parse_comparison() → Parse perbandingan')
    
    add_heading(doc, '4.4 Abstract Syntax Tree (AST)', 2)
    
    # Pengertian AST
    add_heading(doc, 'Pengertian AST', 3)
    doc.add_paragraph(
        'Abstract Syntax Tree (AST) adalah representasi pohon dari struktur sintaksis program. '
        'Berbeda dengan parse tree, AST hanya menyimpan informasi yang relevan dan menghilangkan '
        'detail sintaksis seperti kurung dan koma. AST digunakan untuk analisis semantik dan code generation.'
    )
    
    doc.add_paragraph('Struktur AST CSV_QL:')
    add_code_block(doc, '''SelectStatement
├── columns: List[str]       # ["nama", "nilai"] atau ["*"]
├── table: str               # "data_nilai.csv"
├── where_clause: Optional[Expr]
│   └── BinaryOp
│       ├── left: Expr
│       ├── op: Op
│       └── right: Expr
└── limit: Optional[int]     # 10''')
    
    doc.add_page_break()
    
    # ==================== BAB V ANALISIS SEMANTIK ====================
    add_heading(doc, 'BAB V - ANALISIS SEMANTIK', 1)
    
    # Pengertian Semantic Analysis
    add_heading(doc, 'Pengertian Semantic Analysis', 3)
    doc.add_paragraph(
        'Semantic Analysis (Analisis Semantik) adalah tahap yang memeriksa kebenaran makna program '
        'setelah parsing. Analisis semantik memastikan bahwa program tidak hanya benar secara sintaksis '
        'tetapi juga bermakna. Contoh: memastikan variabel sudah dideklarasikan sebelum digunakan, '
        'tipe data cocok, dll.'
    )
    
    add_heading(doc, '5.1 Fungsi Analisis Semantik', 2)
    doc.add_paragraph('Semantic analyzer CSV_QL melakukan validasi:')
    doc.add_paragraph('1. Validasi File CSV - Memeriksa apakah file yang direferensikan ada')
    doc.add_paragraph('2. Validasi Kolom SELECT - Memeriksa apakah kolom yang diminta ada dalam file')
    doc.add_paragraph('3. Validasi Kolom WHERE - Memeriksa apakah kolom dalam kondisi WHERE ada')
    doc.add_paragraph('4. Validasi LIMIT - Memeriksa apakah nilai LIMIT positif')
    
    add_heading(doc, '5.2 Tabel Simbol', 2)
    
    # Pengertian Symbol Table
    add_heading(doc, 'Pengertian Symbol Table', 3)
    doc.add_paragraph(
        'Symbol Table (Tabel Simbol) adalah struktur data yang menyimpan informasi tentang identifier '
        'dalam program. Informasi ini meliputi nama, tipe, scope, dan atribut lainnya. '
        'Tabel simbol digunakan selama kompilasi untuk pengecekan semantik.'
    )
    
    doc.add_paragraph('Semantic analyzer membangun tabel simbol dari header CSV:')
    add_table(doc,
        ['Symbol', 'Type', 'Source'],
        [
            ['nim', 'String', 'CSV Header'],
            ['nama', 'String', 'CSV Header'],
            ['mata_kuliah', 'String', 'CSV Header'],
            ['sks', 'Number', 'CSV Header'],
            ['nilai_huruf', 'String', 'CSV Header'],
            ['nilai_angka', 'Number', 'CSV Header'],
            ['semester', 'Number', 'CSV Header'],
            ['status', 'String', 'CSV Header'],
        ]
    )
    
    add_heading(doc, '5.3 Contoh Error Handling', 2)
    doc.add_paragraph('Query dengan Error:')
    add_code_block(doc, 'SELECT ipk FROM data_nilai.csv')
    doc.add_paragraph('Output:')
    add_code_block(doc, '❌ Semantic Error: Kolom \'ipk\' tidak ada di data_nilai.csv')
    
    doc.add_page_break()
    
    # ==================== BAB VI IR ====================
    add_heading(doc, 'BAB VI - INTERMEDIATE REPRESENTATION (IR)', 1)
    
    # Pengertian IR
    add_heading(doc, 'Pengertian Intermediate Representation', 3)
    doc.add_paragraph(
        'Intermediate Representation (IR) adalah representasi antara kode sumber dan kode mesin. '
        'IR memudahkan optimasi dan memisahkan front-end (parsing) dari back-end (code generation). '
        'Dalam konteks query engine, IR berupa Query Plan yang menggambarkan langkah-langkah eksekusi.'
    )
    
    add_heading(doc, '6.1 Desain Query Plan', 2)
    add_table(doc,
        ['Step', 'Nama', 'Deskripsi'],
        [
            ['1', 'SCAN', 'Baca file CSV'],
            ['2', 'FILTER', 'Filter baris dengan kondisi WHERE'],
            ['3', 'PROJECT', 'Pilih kolom yang diminta'],
            ['4', 'LIMIT', 'Batasi jumlah hasil'],
        ]
    )
    
    add_heading(doc, '6.2 Struktur Data IR', 2)
    add_code_block(doc, '''@dataclass
class ScanStep:
    table: str                 # Nama file CSV

@dataclass
class FilterStep:
    condition: str             # String representasi kondisi

@dataclass
class ProjectStep:
    columns: List[str]         # Daftar kolom

@dataclass
class LimitStep:
    count: int                 # Jumlah batas

@dataclass
class QueryPlan:
    steps: List[PlanStep]      # Urutan langkah eksekusi''')
    
    add_heading(doc, '6.3 Contoh Query Plan', 2)
    doc.add_paragraph('Query:')
    add_code_block(doc, 'SELECT nama, nilai_huruf FROM data_nilai.csv WHERE status = "Lulus" LIMIT 5')
    
    doc.add_paragraph('Query Plan yang dihasilkan:')
    add_code_block(doc, '''Step 1: SCAN
   └─ Table: data_nilai.csv
Step 2: FILTER
   └─ Condition: status = "Lulus"
Step 3: PROJECT
   └─ Columns: [nama, nilai_huruf]
Step 4: LIMIT
   └─ Count: 5''')
    
    doc.add_page_break()
    
    # ==================== BAB VII EXECUTION ENGINE ====================
    add_heading(doc, 'BAB VII - EXECUTION ENGINE', 1)
    
    # Pengertian Execution Engine
    add_heading(doc, 'Pengertian Execution Engine', 3)
    doc.add_paragraph(
        'Execution Engine adalah komponen yang menjalankan query plan terhadap data aktual. '
        'Engine membaca data dari sumber (file CSV), menerapkan filter, memproyeksikan kolom, '
        'dan mengembalikan hasil. Ini adalah tahap terakhir dari pipeline kompilasi.'
    )
    
    add_heading(doc, '7.1 Arsitektur Pipeline Kompilasi', 2)
    add_code_block(doc, '''Query String
     ↓
┌─────────┐   ┌─────────┐   ┌─────────┐   ┌─────────┐
│  LEXER  │ → │ PARSER  │ → │SEMANTIC │ → │   IR    │
└─────────┘   └─────────┘   └─────────┘   └─────────┘
     ↓             ↓             ↓             ↓
  Tokens         AST         Validated      QueryPlan
                                               ↓
                                         ┌─────────┐
                                         │ ENGINE  │
                                         └─────────┘
                                               ↓
                                          Result Set''')
    
    add_heading(doc, '7.2 Algoritma Eksekusi', 2)
    doc.add_paragraph('Engine mengeksekusi query dengan langkah:')
    doc.add_paragraph('1. Buka dan baca file CSV')
    doc.add_paragraph('2. Untuk setiap baris, evaluasi WHERE clause')
    doc.add_paragraph('3. Jika lolos filter, ambil kolom yang diminta')
    doc.add_paragraph('4. Tambahkan ke hasil, cek LIMIT')
    doc.add_paragraph('5. Kembalikan (headers, rows)')
    
    doc.add_page_break()
    
    # ==================== BAB VIII SIMULASI DFA ====================
    add_heading(doc, 'BAB VIII - SIMULASI AUTOMATA (DFA)', 1)
    
    add_heading(doc, '8.1 Definisi Formal DFA Lexer', 2)
    doc.add_paragraph('DFA M = (Q, Σ, δ, q0, F) dengan:')
    doc.add_paragraph('• Q = {q0, q1, q2, q3, q4, qf} — Himpunan state')
    doc.add_paragraph('• Σ = ASCII characters')
    doc.add_paragraph('• δ = Fungsi transisi (lihat tabel sebelumnya)')
    doc.add_paragraph('• q0 = State awal (START)')
    doc.add_paragraph('• F = {qf} — State akhir (ACCEPT)')
    
    add_heading(doc, '8.2 Contoh Simulasi DFA', 2)
    doc.add_paragraph('Input: SELECT nama FROM data.csv')
    doc.add_paragraph()
    doc.add_paragraph('Trace Transisi:')
    
    add_table(doc,
        ['Step', 'State', 'Input', 'Next State', 'Token'],
        [
            ['1', 'q0', 'S', 'q1', '-'],
            ['2', 'q1', 'E', 'q1', '-'],
            ['3', 'q1', 'L', 'q1', '-'],
            ['4', 'q1', 'E', 'q1', '-'],
            ['5', 'q1', 'C', 'q1', '-'],
            ['6', 'q1', 'T', 'q1', '-'],
            ['7', 'q1', ' ', 'qf', 'SELECT'],
            ['8', 'q0', 'n', 'q1', '-'],
            ['9', 'q1', 'a', 'q1', '-'],
            ['10', 'q1', 'm', 'q1', '-'],
            ['11', 'q1', 'a', 'q1', '-'],
            ['12', 'q1', ' ', 'qf', 'IDENTIFIER(nama)'],
        ]
    )
    
    doc.add_page_break()
    
    # ==================== BAB IX PENGUJIAN ====================
    add_heading(doc, 'BAB IX - PENGUJIAN DAN HASIL', 1)
    
    test_cases = [
        ('Test Case 1: SELECT Semua Data',
         'SELECT * FROM data_nilai.csv',
         '22 baris ditemukan'),
        ('Test Case 2: SELECT Kolom Tertentu',
         'SELECT nim, nama, nilai_huruf FROM data_nilai.csv',
         '22 baris, 3 kolom'),
        ('Test Case 3: Filter WHERE nilai A',
         'SELECT nama, mata_kuliah FROM data_nilai.csv WHERE nilai_huruf = "A"',
         '7 baris ditemukan'),
        ('Test Case 4: Mahasiswa Tidak Lulus',
         'SELECT nim, nama FROM data_nilai.csv WHERE status = "Tidak Lulus"',
         '3 baris ditemukan'),
        ('Test Case 5: Kombinasi AND',
         'SELECT * FROM data_nilai.csv WHERE semester = 5 AND nilai_huruf = "A"',
         '3 baris ditemukan'),
        ('Test Case 6: LIMIT',
         'SELECT * FROM data_nilai.csv LIMIT 5',
         '5 baris ditemukan'),
    ]
    
    for title, query, result in test_cases:
        add_heading(doc, title, 2)
        doc.add_paragraph('Query:')
        add_code_block(doc, query)
        doc.add_paragraph(f'Hasil: {result}')
        doc.add_paragraph()
    
    add_heading(doc, 'Test Case Error Handling', 2)
    doc.add_paragraph('Query dengan kolom tidak ada:')
    add_code_block(doc, 'SELECT ipk FROM data_nilai.csv')
    doc.add_paragraph('Output:')
    add_code_block(doc, '❌ Semantic Error: Kolom \'ipk\' tidak ada di data_nilai.csv')
    
    doc.add_page_break()
    
    # ==================== BAB X KESIMPULAN ====================
    add_heading(doc, 'BAB X - KESIMPULAN', 1)
    
    add_heading(doc, '10.1 Kesimpulan', 2)
    doc.add_paragraph(
        'Proyek CSV_QL berhasil mengimplementasikan mini query engine untuk data nilai mahasiswa '
        'dengan menerapkan konsep automata dan teknik kompilasi:'
    )
    doc.add_paragraph('1. Lexical Analyzer - Diimplementasikan menggunakan DFA untuk tokenisasi query SQL')
    doc.add_paragraph('2. Parser - Diimplementasikan menggunakan metode Recursive Descent dengan CFG')
    doc.add_paragraph('3. Semantic Analyzer - Melakukan validasi file CSV dan kolom')
    doc.add_paragraph('4. Intermediate Representation - Query plan (SCAN, FILTER, PROJECT, LIMIT)')
    doc.add_paragraph('5. Execution Engine - Mengeksekusi query terhadap file CSV')
    
    add_heading(doc, '10.2 Saran Pengembangan', 2)
    doc.add_paragraph('1. Mendukung klausa ORDER BY, GROUP BY, dan HAVING')
    doc.add_paragraph('2. Mendukung fungsi agregat (COUNT, SUM, AVG)')
    doc.add_paragraph('3. Mendukung JOIN antar file CSV')
    doc.add_paragraph('4. Optimasi query plan')
    
    doc.add_page_break()
    
    # ==================== DAFTAR PUSTAKA ====================
    add_heading(doc, 'DAFTAR PUSTAKA', 1)
    
    references = [
        'Aho, A. V., Lam, M. S., Sethi, R., & Ullman, J. D. (2006). Compilers: Principles, Techniques, and Tools (2nd ed.). Addison-Wesley.',
        'Hopcroft, J. E., Motwani, R., & Ullman, J. D. (2006). Introduction to Automata Theory, Languages, and Computation (3rd ed.). Pearson.',
        'Grune, D., & Jacobs, C. J. (2007). Parsing Techniques: A Practical Guide (2nd ed.). Springer.',
        'Python Software Foundation. (2024). Python CSV Module Documentation. https://docs.python.org/3/library/csv.html',
        'Garcia-Molina, H., Ullman, J. D., & Widom, J. (2008). Database Systems: The Complete Book (2nd ed.). Pearson.',
    ]
    
    for i, ref in enumerate(references, 1):
        doc.add_paragraph(f'[{i}] {ref}')
    
    doc.add_page_break()
    
    # ==================== LAMPIRAN ====================
    add_heading(doc, 'LAMPIRAN', 1)
    
    add_heading(doc, 'A. Struktur Direktori Proyek', 2)
    add_code_block(doc, '''csv_ql/
├── README.md              Dokumentasi proyek
├── TEST_CASES.md          Test cases
├── MAKALAH_CSV_QL.docx    Makalah ini
├── data_nilai.csv         Data sampel
│
└── src/                   Source code
    ├── main.py            Entry point & REPL
    ├── tokens.py          Definisi token
    ├── lexer.py           Lexical analyzer (DFA)
    ├── ast_nodes.py       Abstract Syntax Tree
    ├── parser.py          Syntax analyzer (CFG)
    ├── semantic.py        Semantic analyzer
    ├── ir.py              Intermediate representation
    ├── engine.py          Query execution
    └── dfa.py             DFA visualization''')
    
    add_heading(doc, 'B. Cara Menjalankan Program', 2)
    add_code_block(doc, '''cd csv_ql/src

# Mode REPL (Interactive)
python main.py

# Mode Direct Query
python main.py "SELECT * FROM ../data_nilai.csv"

# Mode Verbose (dengan detail kompilasi)
python main.py "SELECT * FROM ../data_nilai.csv" -v''')
    
    add_heading(doc, 'C. Screenshot Aplikasi', 2)
    doc.add_paragraph('[Tambahkan screenshot hasil running program di sini]')
    
    return doc

if __name__ == '__main__':
    doc = create_document()
    doc.save('MAKALAH_CSV_QL.docx')
    print('✅ MAKALAH_CSV_QL.docx berhasil dibuat!')
